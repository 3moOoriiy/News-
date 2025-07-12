import streamlit as st
import feedparser
import pandas as pd
from datetime import datetime, timedelta
from io import BytesIO
from textblob import TextBlob
from collections import Counter
from docx import Document
import urllib.request
import urllib.parse
import json
import re
import time
from bs4 import BeautifulSoup
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
import requests

st.set_page_config(page_title=":newspaper: أداة الأخبار العربية الذكية", layout="wide")
st.title(":rolled_up_newspaper: أداة إدارة وتحليل الأخبار المتطورة (RSS + Web Scraping)")

# التصنيفات المحسّنة
category_keywords = input("please insert keywords:")
# الدوال المحسّنة
def summarize(text, max_words=30):
    if not text:
        return "لا يوجد ملخص متاح"
    words = text.split()
    if len(words) <= max_words:
        return text
    return " ".join(words[:max_words]) + "..."

def analyze_sentiment(text):
    if not text:
        return ":neutral_face: محايد"
    try:
        polarity = TextBlob(text).sentiment.polarity
        if polarity > 0.1:
            return ":smiley: إيجابي"
        elif polarity < -0.1:
            return ":angry: سلبي"
        else:
            return ":neutral_face: محايد"
    except:
        return ":neutral_face: محايد"

def detect_category(text):
    if not text:
        return "غير مصنّف"
    text_lower = text.lower()
    category_scores = {}
    
    for category, words in category_keywords.items():
        score = sum(1 for word in words if word in text_lower)
        if score > 0:
            category_scores[category] = score
    
    if category_scores:
        return max(category_scores, key=category_scores.get)
    return "غير مصنّف"

def safe_request(url, timeout=10):
    """طلب آمن مع معالجة الأخطاء"""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        req = urllib.request.Request(url, headers=headers)
        response = urllib.request.urlopen(req, timeout=timeout)
        return response.read().decode('utf-8', errors='ignore')
    except Exception as e:
        st.warning(f"خطأ في الوصول لـ {url}: {str(e)}")
        return None

def fetch_multiple_pages(base_url, max_pages=5):
    """جلب محتوى من عدة صفحات"""
    all_html = []
    for page in range(1, max_pages + 1):
        try:
            # تعديل الرابط لإضافة رقم الصفحة
            if "?" in base_url:
                page_url = f"{base_url}&page={page}"
            else:
                page_url = f"{base_url}?page={page}"
            
            html = safe_request(page_url)
            if html:
                all_html.append(html)
                time.sleep(1)  # تجنب حظر IP
        except:
            continue
    return all_html

def get_dynamic_page(url):
    """جلب محتوى الصفحات الديناميكية باستخدام Selenium"""
    try:
        options = Options()
        options.headless = True
        driver = webdriver.Chrome(options=options)
        driver.get(url)
        time.sleep(3)  # انتظار تحميل المحتوى
        html = driver.page_source
        driver.quit()
        return html
    except Exception as e:
        st.error(f"خطأ في جلب الصفحة الديناميكية: {str(e)}")
        return None

def fetch_from_api(api_url):
    """جلب البيانات من واجهات API"""
    try:
        response = requests.get(api_url, headers={'User-Agent': 'Mozilla/5.0'})
        return response.json()
    except Exception as e:
        st.error(f"خطأ في جلب البيانات من API: {str(e)}")
        return None

def parse_with_bs4(html):
    """استخراج محتوى متقدم باستخدام BeautifulSoup"""
    try:
        soup = BeautifulSoup(html, 'html.parser')
        articles = soup.find_all('article')  # أو أي علامة أخرى يستخدمها الموقع
        
        news_list = []
        for article in articles:
            title = article.find('h2').text if article.find('h2') else ""
            link = article.find('a')['href'] if article.find('a') else ""
            summary = article.find('p').text if article.find('p') else title
            
            news_list.append({
                'title': title.strip(),
                'summary': summarize(summary.strip()),
                'link': link,
                'published': datetime.now(),
                'image': "",
                'sentiment': analyze_sentiment(title),
                'category': detect_category(title),
                'extraction_method': 'BeautifulSoup'
            })
        return news_list
    except Exception as e:
        st.error(f"خطأ في تحليل المحتوى: {str(e)}")
        return []

def extract_news_from_html(html_content, source_name, base_url):
    """استخراج الأخبار من HTML بطريقة ذكية"""
    if not html_content:
        return []
    
    news_list = []
    
    # البحث عن العناوين المحتملة
    title_patterns = [
        r'<h[1-4][^>]*>(.*?)</h[1-4]>',
        r'<title[^>]*>(.*?)</title>',
        r'<a[^>]*title="([^"]+)"',
        r'<div[^>]*class="[^"]*title[^"]*"[^>]*>(.*?)</div>'
    ]
    
    # البحث عن الروابط
    link_patterns = [
        r'<a[^>]*href="([^"]+)"[^>]*>(.*?)</a>',
        r'href="([^"]+)"'
    ]
    
    titles = []
    links = []
    
    for pattern in title_patterns:
        matches = re.findall(pattern, html_content, re.IGNORECASE | re.DOTALL)
        for match in matches:
            if isinstance(match, tuple):
                title = match[0] if match[0] else match[1] if len(match) > 1 else ""
            else:
                title = match
            title = re.sub(r'<[^>]+>', '', title).strip()
            if title and len(title) > 10 and len(title) < 200:
                titles.append(title)
    
    for pattern in link_patterns:
        matches = re.findall(pattern, html_content, re.IGNORECASE)
        for match in matches:
            if isinstance(match, tuple):
                link = match[0]
            else:
                link = match
            if link and not link.startswith('#') and not link.startswith('javascript:'):
                if link.startswith('/'):
                    link = base_url + link
                elif not link.startswith('http'):
                    link = base_url + '/' + link
                links.append(link)
    
    # دمج العناوين والروابط
    for i, title in enumerate(titles[:50]):  # أول 50 خبر
        link = links[i] if i < len(links) else base_url
        
        news_list.append({
            "source": source_name,
            "title": title,
            "summary": title,  # استخدام العنوان كملخص مؤقت
            "link": link,
            "published": datetime.now(),
            "image": "",
            "sentiment": analyze_sentiment(title),
            "category": detect_category(title),
            "extraction_method": "HTML Parsing"
        })
    
    return news_list

def fetch_rss_news(source_name, url, keywords, date_from, date_to, chosen_category):
    """إصدارة محسنة مع إصلاح فلترة التاريخ"""
    try:
        feed = feedparser.parse(url)
        news_list = []
        
        if not hasattr(feed, 'entries') or len(feed.entries) == 0:
            return []
        
        for entry in feed.entries:
            try:
                title = entry.get('title', 'بدون عنوان')
                summary = entry.get('summary', entry.get('description', title))
                link = entry.get('link', '')
                published = entry.get('published', '')
                
                # معالجة التاريخ بشكل محسن
                published_dt = None
                try:
                    published_dt = datetime.strptime(published, "%a, %d %b %Y %H:%M:%S %Z")
                except:
                    try:
                        published_dt = datetime.strptime(published, "%Y-%m-%dT%H:%M:%S%z")
                    except:
                        try:
                            published_dt = datetime.strptime(published, "%Y-%m-%d %H:%M:%S")
                        except:
                            published_dt = datetime.now()
                
                # تحويل إلى تاريخ بدون وقت للمقارنة
                published_date = published_dt.date()
                
                # فلترة التاريخ بشكل صحيح
                if not (date_from <= published_date <= date_to):
                    continue

                # فلترة الكلمات المفتاحية
                full_text = title + " " + summary
                if keywords:
                    if isinstance(keywords, str):
                        keywords = [k.strip() for k in keywords.split(",") if k.strip()]
                    
                    if not any(re.search(r'\b{}\b'.format(re.escape(k.lower())), full_text.lower()) for k in keywords):
                        continue

                # فلترة التصنيف
                auto_category = detect_category(full_text)
                if chosen_category != "الكل" and auto_category != chosen_category:
                    continue

                # البحث عن صورة
                image = ""
                if hasattr(entry, 'media_content') and entry.media_content:
                    image = entry.media_content[0].get('url', '')
                elif hasattr(entry, 'media_thumbnail') and entry.media_thumbnail:
                    image = entry.media_thumbnail[0].get('url', '')

                news_list.append({
                    "source": source_name,
                    "title": title,
                    "summary": summary,
                    "link": link,
                    "published": published_dt,
                    "image": image,
                    "sentiment": analyze_sentiment(summary),
                    "category": auto_category,
                    "extraction_method": "RSS"
                })
                
            except Exception as e:
                continue
                
        return news_list
        
    except Exception as e:
        st.error(f"خطأ في جلب أخبار RSS: {str(e)}")
        return []

def fetch_website_news(source_name, url, keywords, date_from, date_to, chosen_category, max_pages=5, method="auto"):
    """إصدارة محسنة مع زيادة عدد الصفحات"""
    try:
        st.info(f":arrows_counterclockwise: جاري تحليل موقع {source_name}...")
        
        # جلب محتوى من عدة صفحات (زيادة عدد الصفحات إلى 5)
        all_html = []
        if method == "dynamic":
            html_content = get_dynamic_page(url)
            if html_content:
                all_html.append(html_content)
        elif method == "api" and "api_url" in iraqi_news_sources.get(source_name, {}):
            api_data = fetch_from_api(iraqi_news_sources[source_name]["api_url"])
            return process_api_data(api_data, source_name, keywords, date_from, date_to, chosen_category)
        else:
            all_html = fetch_multiple_pages(url, max_pages)
        
        if not all_html:
            return []
        
        # استخراج الأخبار من HTML
        base_url = url.rstrip('/')
        news_list = []
        
        for html in all_html:
            if method == "bs4":
                news_list.extend(parse_with_bs4(html))
            else:
                news_list.extend(extract_news_from_html(html, source_name, base_url))
        
        # فلترة النتائج
        filtered_news = []
        for news in news_list:
            # لا نطبق فلترة التاريخ على الأخبار من المواقع مباشرة
            # لأنها عادة لا تحتوي على تواريخ دقيقة
            
            # فلترة الكلمات المفتاحية
            full_text = news['title'] + " " + news['summary']
            if keywords:
                if isinstance(keywords, str):
                    keywords = [k.strip() for k in keywords.split(",") if k.strip()]
                
                if not any(re.search(r'\b{}\b'.format(re.escape(k.lower())), full_text.lower()) for k in keywords):
                    continue
            
            # فلترة التصنيف
            if chosen_category != "الكل" and news['category'] != chosen_category:
                continue
            
            filtered_news.append(news)
        
        return filtered_news[:50]  # زيادة الحد إلى 50 خبر
        
    except Exception as e:
        st.error(f"خطأ في جلب الأخبار من {source_name}: {str(e)}")
        return []

def smart_news_fetcher(source_name, source_info, keywords, date_from, date_to, chosen_category, method="auto", max_pages=5):
    """جالب الأخبار الذكي - يجرب عدة طرق"""
    all_news = []
    
    # المحاولة الأولى: RSS
    if method in ["auto", "rss"] and source_info.get("rss_options"):
        st.info(":arrows_counterclockwise: المحاولة الأولى: البحث عن RSS...")
        for rss_url in source_info["rss_options"]:
            try:
                news = fetch_rss_news(source_name, rss_url, keywords, date_from, date_to, chosen_category)
                if news:
                    st.success(f":white_check_mark: تم العثور على {len(news)} خبر من RSS: {rss_url}")
                    all_news.extend(news)
                    if method == "rss":
                        return all_news  # إذا كان الخيار RSS فقط
                    break
            except:
                continue
    
    # المحاولة الثانية: تحليل الموقع مباشرة
    if method in ["auto", "html", "dynamic", "bs4", "api"]:
        st.info(":arrows_counterclockwise: المحاولة الثانية: تحليل الموقع مباشرة...")
        website_news = fetch_website_news(
            source_name,
            source_info["url"],
            keywords,
            date_from,
            date_to,
            chosen_category,
            max_pages,
            method if method != "auto" else "html"
        )
        if website_news:
            st.success(f":white_check_mark: تم استخراج {len(website_news)} خبر من الموقع مباشرة")
            all_news.extend(website_news)
    
    # إزالة المكرر
    seen_titles = set()
    unique_news = []
    for news in all_news:
        if news['title'] not in seen_titles:
            seen_titles.add(news['title'])
            unique_news.append(news)
    
    return unique_news

def export_to_word(news_list):
    doc = Document()
    doc.add_heading('تقرير الأخبار المجمعة', 0)
    doc.add_paragraph(f'تاريخ التقرير: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'عدد الأخبار: {len(news_list)}')
    doc.add_paragraph('---')
    
    for i, news in enumerate(news_list, 1):
        doc.add_heading(f'{i}. {news["title"]}', level=2)
        doc.add_paragraph(f"المصدر: {news['source']}")
        doc.add_paragraph(f"التصنيف: {news['category']}")
        doc.add_paragraph(f"التاريخ: {news['published'].strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"طريقة الاستخراج: {news.get('extraction_method', 'غير محدد')}")
        doc.add_paragraph(f"التحليل العاطفي: {news['sentiment']}")
        doc.add_paragraph(f"الملخص: {news['summary']}")
        doc.add_paragraph(f"الرابط: {news['link']}")
        doc.add_paragraph('---')
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def export_to_excel(news_list):
    df = pd.DataFrame(news_list)
    columns_order = ['source', 'title', 'category', 'sentiment', 'published', 'summary', 'link', 'extraction_method']
    df = df.reindex(columns=[col for col in columns_order if col in df.columns])
    
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='الأخبار')
    buffer.seek(0)
    return buffer

def process_api_data(api_data, source_name, keywords, date_from, date_to, chosen_category):
    """معالجة بيانات API"""
    if not api_data:
        return []
    
    news_list = []
    for item in api_data:
        try:
            title = item.get('title', '')
            summary = item.get('summary', title)
            link = item.get('url', '')
            published = item.get('published', '')
            
            # معالجة التاريخ
            try:
                published_dt = datetime.strptime(published, "%Y-%m-%dT%H:%M:%SZ")
            except:
                try:
                    published_dt = datetime.strptime(published, "%Y-%m-%d %H:%M:%S")
                except:
                    published_dt = datetime.now()
            
            # فلترة التاريخ
            if not (date_from <= published_dt.date() <= date_to):
                continue

            # فلترة الكلمات المفتاحية
            full_text = title + " " + summary
            if keywords:
                if isinstance(keywords, str):
                    keywords = [k.strip() for k in keywords.split(",") if k.strip()]
                
                if not any(re.search(r'\b{}\b'.format(re.escape(k.lower())), full_text.lower()) for k in keywords):
                    continue

            # فلترة التصنيف
            auto_category = detect_category(full_text)
            if chosen_category != "الكل" and auto_category != chosen_category:
                continue

            news_list.append({
                "source": source_name,
                "title": title,
                "summary": summary,
                "link": link,
                "published": published_dt,
                "image": item.get('image', ''),
                "sentiment": analyze_sentiment(summary),
                "category": auto_category,
                "extraction_method": "API"
            })
            
        except Exception as e:
            continue
    
    return news_list

# مصادر الأخبار المحسّنة مع إضافة واجهات API
general_rss_feeds = {
    "BBC عربي": "http://feeds.bbci.co.uk/arabic/rss.xml",
    "الجزيرة": "https://www.aljazeera.net/aljazeerarss/ar/home",
    "RT Arabic": "https://arabic.rt.com/rss/",
    "France24 عربي": "https://www.france24.com/ar/rss",
    "سكاي نيوز عربية": "https://www.skynewsarabia.com/web/rss",
    "عربي21": "https://arabi21.com/feed"
}

iraqi_news_sources = {
    "وزارة الداخلية العراقية": {
        "url": "https://moi.gov.iq/",
        "type": "website",
        "rss_options": [
            "https://moi.gov.iq/feed/",
            "https://moi.gov.iq/rss.xml"
        ],
        "api_url": "https://moi.gov.iq/api/news"
    },
    "هذا اليوم": {
        "url": "https://hathalyoum.net/",
        "type": "website",
        "rss_options": [
            "https://hathalyoum.net/feed/",
            "https://hathalyoum.net/rss.xml"
        ]
    },
    "العراق اليوم": {
        "url": "https://iraqtoday.com/",
        "type": "website",
        "rss_options": [
            "https://iraqtoday.com/feed/",
            "https://iraqtoday.com/rss.xml"
        ]
    },
    "رئاسة الجمهورية العراقية": {
        "url": "https://presidency.iq/default.aspx",
        "type": "website",
        "rss_options": [
            "https://presidency.iq/feed/",
            "https://presidency.iq/rss.xml"
        ]
    },
    "الشرق الأوسط": {
        "url": "https://asharq.com/",
        "type": "website",
        "rss_options": [
            "https://asharq.com/feed/",
            "https://asharq.com/rss.xml"
        ]
    },
    "RT Arabic - العراق": {
        "url": "https://arabic.rt.com/focuses/10744-%D8%A7%D9%84%D8%B9%D8%B1%D8%A7%D9%82/",
        "type": "website",
        "rss_options": [
            "https://arabic.rt.com/rss/"
        ]
    },
    "إندبندنت عربية": {
        "url": "https://www.independentarabia.com/",
        "type": "website",
        "rss_options": [
            "https://www.independentarabia.com/rss"
        ]
    },
    "فرانس 24 عربي": {
        "url": "https://www.france24.com/ar/",
        "type": "website",
        "rss_options": [
            "https://www.france24.com/ar/rss"
        ]
    }
}

world_news_sources = {
    "CNN عربي": {
        "url": "https://arabic.cnn.com/",
        "type": "website",
        "rss_options": [
            "https://arabic.cnn.com/feed/",
            "https://arabic.cnn.com/rss.xml"
        ]
    },
    "Axios": {
        "url": "https://www.axios.com/",
        "type": "website",
        "rss_options": [
            "https://api.axios.com/feed/",
            "https://www.axios.com/feeds/feed.xml"
        ]
    },
    "BBC News": {
        "url": "https://www.bbc.com/news",
        "type": "website",
        "rss_options": [
            "http://feeds.bbci.co.uk/news/rss.xml",
            "https://feeds.bbci.co.uk/news/world/rss.xml"
        ]
    },
    "i24NEWS عربي": {
        "url": "https://www.i24news.tv/ar",
        "type": "website",
        "rss_options": [
            "https://www.i24news.tv/ar/rss",
            "https://www.i24news.tv/ar/feed/"
        ]
    },
    "France24 إنجليزي": {
        "url": "https://www.france24.com/en/",
        "type": "website",
        "rss_options": [
            "https://www.france24.com/en/rss",
            "https://www.france24.com/en/africa/rss"
        ]
    },
    "SwissInfo عربي": {
        "url": "https://www.swissinfo.ch/ara/",
        "type": "website",
        "rss_options": [
            "https://www.swissinfo.ch/ara/rss",
            "https://www.swissinfo.ch/~rss/ara"
        ]
    },
    "Reuters": {
        "url": "https://www.reuters.com/",
        "type": "website",
        "rss_options": [
            "https://feeds.reuters.com/reuters/topNews",
            "https://feeds.reuters.com/Reuters/worldNews"
        ]
    },
    "AP News": {
        "url": "https://apnews.com/",
        "type": "website",
        "rss_options": [
            "https://feeds.apnews.com/rss/apf-topnews",
            "https://feeds.apnews.com/rss/apf-intlnews"
        ]
    },
    "NBC News": {
        "url": "https://www.nbcnews.com/",
        "type": "website",
        "rss_options": [
            "https://feeds.nbcnews.com/nbcnews/public/news",
            "https://feeds.nbcnews.com/nbcnews/public/world"
        ]
    },
    "ABC News": {
        "url": "https://abcnews.go.com/",
        "type": "website",
        "rss_options": [
            "https://abcnews.go.com/abcnews/topstories",
            "https://abcnews.go.com/abcnews/internationalheadlines"
        ]
    },
    "The Independent": {
        "url": "https://www.independent.co.uk/",
        "type": "website",
        "rss_options": [
            "https://www.independent.co.uk/rss",
            "https://www.independent.co.uk/news/world/rss"
        ]
    },
    "RT Arabic العالمي": {
        "url": "https://arabic.rt.com/",
        "type": "website",
        "rss_options": [
            "https://arabic.rt.com/rss/",
            "https://arabic.rt.com/rss/world/"
        ]
    },
    "Sky News العربية العالمي": {
        "url": "https://sarabic.ae/",
        "type": "website",
        "rss_options": [
            "https://sarabic.ae/feed/",
            "https://sarabic.ae/rss.xml"
        ]
    }
}

# واجهة المستخدم المحسّنة
st.sidebar.header(":gear: إعدادات البحث المتقدم")

# اختيار نوع المصدر
source_type = st.sidebar.selectbox(
    ":earth_africa: اختر نوع المصدر:",
    ["المصادر العامة", "المصادر العراقية", "أبرز الأخبار في العالم"],
    help="المصادر العامة تعتمد على RSS، المصادر العراقية تستخدم تقنيات متقدمة، وأبرز الأخبار في العالم تغطي المصادر الإخبارية العالمية الرئيسية"
)

# اختيار طريقة الجلب
scraping_method = st.sidebar.selectbox(
    ":mag_right: طريقة الجلب:",
    ["auto", "rss", "html", "dynamic", "bs4", "api"],
    help="اختر طريقة جلب المحتوى: auto (تلقائي), rss (RSS فقط), html (تحليل HTML), dynamic (صفحات ديناميكية), bs4 (BeautifulSoup), api (واجهة برمجة)"
)

if source_type == "المصادر العامة":
    selected_source = st.sidebar.selectbox(":globe_with_meridians: اختر مصدر الأخبار:", list(general_rss_feeds.keys()))
    source_url = general_rss_feeds[selected_source]
    source_info = {"type": "rss", "url": source_url}
elif source_type == "المصادر العراقية":
    selected_source = st.sidebar.selectbox(":flag-iq: اختر مصدر الأخبار العراقي:", list(iraqi_news_sources.keys()))
    source_info = iraqi_news_sources[selected_source]
else:  # أبرز الأخبار في العالم
    selected_source = st.sidebar.selectbox(":earth_americas: اختر مصدر الأخبار العالمي:", list(world_news_sources.keys()))
    source_info = world_news_sources[selected_source]

# إعدادات البحث
keywords_input = st.sidebar.text_input(
    ":mag: كلمات مفتاحية (مفصولة بفواصل):", 
    "",
    help="يمكنك إدخال أي كلمات تريد البحث عنها"
)
keywords = keywords_input

category_filter = st.sidebar.selectbox(
    ":file_folder: اختر التصنيف:", 
    ["الكل"] + list(category_keywords.keys()),
    help="فلترة الأخبار حسب التصنيف"
)

# إعدادات التاريخ
col_date1, col_date2 = st.sidebar.columns(2)
with col_date1:
    date_from = st.date_input(":date: من تاريخ:", datetime.today() - timedelta(days=14))
with col_date2:
    date_to = st.date_input(":date: إلى تاريخ:", datetime.today())

# خيارات متقدمة
with st.sidebar.expander(":gear: خيارات متقدمة"):
    max_news = st.slider("عدد الأخبار الأقصى:", 5, 100, 50)  # زيادة الحد الأقصى
    max_pages = st.slider("عدد الصفحات للبحث:", 1, 10, 5)  # زيادة عدد الصفحات الافتراضي
    include_sentiment = st.checkbox("تحليل المشاعر", True)
    include_categorization = st.checkbox("التصنيف التلقائي", True)
    image_size = st.slider("حجم الصور:", 100, 500, 200)

run = st.sidebar.button(":inbox_tray: جلب الأخبار", type="primary", help="ابدأ عملية جلب وتحليل الأخبار")

# عرض النتائج
if run:
    with st.spinner(":robot_face: جاري تشغيل الذكاء الاصطناعي لجلب الأخبار..."):
        start_time = time.time()
        
        news = smart_news_fetcher(
            selected_source,
            source_info,
            keywords,
            date_from,
            date_to,
            category_filter,
            scraping_method,
            max_pages
        )
        
        end_time = time.time()
        processing_time = round(end_time - start_time, 2)
    
    if news:
        st.success(f":tada: تم جلب {len(news)} خبر من {selected_source} في {processing_time} ثانية")
        
        # إحصائيات سريعة
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric(":newspaper: إجمالي الأخبار", len(news))
        with col2:
            categories = [n['category'] for n in news]
            st.metric(":file_folder: أكثر تصنيف", Counter(categories).most_common(1)[0][0] if categories else "غير محدد")
        with col3:
            positive_news = len([n for n in news if "إيجابي" in n['sentiment']])
            st.metric(":smiley: أخبار إيجابية", positive_news)
        with col4:
            st.metric(":stopwatch: وقت المعالجة", f"{processing_time}s")
        
        # عرض الأخبار
        st.subheader(":bookmark_tabs: الأخبار المجمعة")
        
        for i, item in enumerate(news[:max_news], 1):
            with st.container():
                st.markdown(f"### {i}. :newspaper: {item['title']}")
                
                col_info, col_content = st.columns([1, 2])
                
                with col_info:
                    st.markdown(f"**:office: المصدر:** {item['source']}")
                    st.markdown(f"**:date: التاريخ:** {item['published'].strftime('%Y-%m-%d %H:%M')}")
                    st.markdown(f"**:file_folder: التصنيف:** {item['category']}")
                    st.markdown(f"**:performing_arts: المشاعر:** {item['sentiment']}")
                    st.markdown(f"**:wrench: الطريقة:** {item.get('extraction_method', 'غير محدد')}")
                
                with col_content:
                    st.markdown(f"**:page_facing_up: الملخص:** {summarize(item['summary'], 40)}")
                    st.markdown(f"**:link: [قراءة المقال كاملاً ↗]({item['link']})**")
                
                if item.get('image'):
                    st.image(item['image'], caption=item['title'], width=image_size)
                
                st.markdown("---")
        
        # تصدير البيانات
        st.subheader(":outbox_tray: تصدير البيانات")
        col_export1, col_export2, col_export3 = st.columns(3)
        
        with col_export1:
            word_file = export_to_word(news)
            st.download_button(
                ":page_facing_up: تحميل Word",
                data=word_file,
                file_name=f"اخبار_{selected_source}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        
        with col_export2:
            excel_file = export_to_excel(news)
            st.download_button(
                ":bar_chart: تحميل Excel",
                data=excel_file,
                file_name=f"اخبار_{selected_source}_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
        
        with col_export3:
            json_data = json.dumps(news, ensure_ascii=False, default=str, indent=2)
            st.download_button(
                ":floppy_disk: تحميل JSON",
                data=json_data.encode('utf-8'),
                file_name=f"اخبار_{selected_source}_{datetime.now().strftime('%Y%m%d_%H%M')}.json",
                mime="application/json"
            )
        
        # تحليلات متقدمة
        with st.expander(":bar_chart: تحليلات متقدمة"):
            col_analysis1, col_analysis2 = st.columns(2)
            
            with col_analysis1:
                st.subheader(":file_folder: توزيع التصنيفات")
                categories = [n['category'] for n in news]
                category_counts = Counter(categories)
                
                for cat, count in category_counts.most_common():
                    percentage = (count / len(news)) * 100
                    st.write(f"• **{cat}**: {count} ({percentage:.1f}%)")
            
            with col_analysis2:
                st.subheader(":performing_arts: تحليل المشاعر")
                sentiments = [n['sentiment'] for n in news]
                sentiment_counts = Counter(sentiments)
                
                for sent, count in sentiment_counts.items():
                    percentage = (count / len(news)) * 100
                    st.write(f"• **{sent}**: {count} ({percentage:.1f}%)")
            
            st.subheader(":abc: أكثر الكلمات تكراراً")
            all_text = " ".join([n['title'] + " " + n['summary'] for n in news])
            words = re.findall(r'\b[أ-ي]{3,}\b', all_text)
            word_freq = Counter(words).most_common(15)
            
            if word_freq:
                cols = st.columns(3)
                for i, (word, freq) in enumerate(word_freq):
                    with cols[i % 3]:
                        st.write(f"**{word}**: {freq} مرة")
    
    else:
        st.warning(":x: لم يتم العثور على أخبار بالشروط المحددة")
        st.info(":bulb: جرب توسيع نطاق التاريخ أو تغيير الكلمات المفتاحية")
        st.markdown(f":link: **[زيارة {selected_source} مباشرة]({source_info['url']})**")

# معلومات في الشريط الجانبي
st.sidebar.markdown("---")
st.sidebar.info("""
:rocket: **تقنيات متقدمة:**
- جلب RSS تلقائي
- تحليل مواقع الويب
- تصنيف ذكي للأخبار
- تحليل المشاعر
- إزالة المحتوى المكرر
- دعم الصفحات المتعددة
- استخراج من واجهات API
- معالجة الصفحات الديناميكية
""")

st.sidebar.success(":white_check_mark: نظام ذكي متطور لجمع الأخبار!")

# معلومات تقنية
with st.expander(":information_source: معلومات تقنية"):
    st.markdown("""
    ### :hammer_and_wrench: التقنيات المستخدمة:
    - **RSS Parsing**: لجلب الأخبار من المصادر التقليدية
    - **HTML Analysis**: لتحليل مواقع الويب مباشرة  
    - **Dynamic Page Loading**: باستخدام Selenium للصفحات الديناميكية
    - **API Integration**: لجلب البيانات من واجهات برمجة التطبيقات
    - **Multi-Page Crawling**: التنقل عبر صفحات الموقع
    - **Smart Categorization**: تصنيف تلقائي للأخبار
    - **Sentiment Analysis**: تحليل المشاعر باستخدام TextBlob
    
    ### :dart: كيف يعمل النظام:
    1. **محاولة RSS أولاً**: البحث عن feeds متاحة
    2. **تحليل HTML**: استخراج المحتوى من الصفحة مباشرة
    3. **التنقل عبر الصفحات**: جلب الأخبار من صفحات متعددة
    4. **معالجة ذكية**: تنظيف وتصنيف البيانات
    5. **إزالة التكرار**: ضمان عدم تكرار الأخبار
    6. **تحليل متقدم**: استخراج الإحصائيات والمشاعر
    """)
